"""GB/T 9704-2012 国家公文标准格式 Docx 导出引擎。

严格按国家公文标准规范排版：
- 版面：A4 纸张，上 37mm、下 35mm、左 28mm、右 26mm（或按单位自定义模板调整）
- 标题：方正小标宋_GB2312（二号 22pt），居中
- 主送机关：仿宋_GB2312（三号 16pt），顶格
- 一级标题：黑体（三号 16pt），首行缩进 2 字符
- 二级标题：楷体_GB2312（三号 16pt），首行缩进 2 字符
- 正文段落：仿宋_GB2312（三号 16pt），首行缩进 2 字符，固定行距 28 磅
- 成文日期/发文机关：右对齐
"""

import io
import re
from typing import BinaryIO

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Mm, Pt, RGBColor

from ..domain.schemas import Project, TemplateConfig


def _set_cell_border(cell, **kwargs):
    """设置单元格边框。"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:bottom w:val="{kwargs.get("bottom", "none")}" w:sz="{kwargs.get("sz", "4")}" w:space="0" w:color="{kwargs.get("color", "auto")}"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def export_project_to_docx(project: Project, template: TemplateConfig = None) -> io.BytesIO:
    """将项目草稿渲染为符合 GB/T 9704-2012 或单位自定义模板的 Word (.docx) 二进制流。"""
    tpl = template or project.template_config or TemplateConfig()

    doc = docx.Document()

    # 1. 页面设置（A4 纸张 + 标准边距）
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(tpl.top_margin_mm)
    section.bottom_margin = Mm(tpl.bottom_margin_mm)
    section.left_margin = Mm(tpl.left_margin_mm)
    section.right_margin = Mm(tpl.right_margin_mm)

    # 2. 红头/发文机关（若有配置）
    if tpl.header_text:
        p_hdr = doc.add_paragraph()
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_hdr.paragraph_format.space_before = Pt(0)
        p_hdr.paragraph_format.space_after = Pt(12)
        p_hdr.paragraph_format.line_spacing = Pt(36)
        r_hdr = p_hdr.add_run(tpl.header_text)
        r_hdr.font.name = tpl.title_font or "方正小标宋_GB2312"
        r_hdr.font.size = Pt(28)
        r_hdr.font.bold = True
        r_hdr.font.color.rgb = RGBColor(218, 41, 28)  # 官方红头朱红色
        r_hdr._element.rPr.rFonts.set(qn("w:eastAsia"), tpl.title_font or "方正小标宋_GB2312")

        # 发文字号（若有）
        if tpl.doc_code:
            p_code = doc.add_paragraph()
            p_code.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_code.paragraph_format.space_before = Pt(0)
            p_code.paragraph_format.space_after = Pt(6)
            p_code.paragraph_format.line_spacing = Pt(20)
            r_code = p_code.add_run(tpl.doc_code)
            r_code.font.name = "仿宋_GB2312"
            r_code.font.size = Pt(12)
            r_code._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")

        # 分割红线
        p_line = doc.add_paragraph()
        p_line.paragraph_format.space_after = Pt(18)
        p_line_border = parse_xml(
            f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="18" w:space="1" w:color="DA291C"/></w:pBdr>'
        )
        p_line._p.get_or_add_pPr().append(p_line_border)

    # 解析草稿文本
    draft_text = project.draft.strip() if project.draft else f"{project.name}\n\n（暂无草稿内容）"
    lines = [line.strip() for line in draft_text.splitlines() if line.strip()]

    if not lines:
        lines = [project.name, "（正文为空）"]

    # 3. 提取并排版标题（第一行）
    title_line = lines[0]
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(16)
    p_title.paragraph_format.line_spacing = Pt(30)
    r_title = p_title.add_run(title_line)
    r_title.font.name = tpl.title_font or "方正小标宋_GB2312"
    r_title.font.size = Pt(tpl.title_size_pt or 22.0)
    r_title.font.bold = True
    r_title._element.rPr.rFonts.set(qn("w:eastAsia"), tpl.title_font or "方正小标宋_GB2312")

    # 4. 正文段落排版
    for line in lines[1:]:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(tpl.line_spacing_pt or 28.0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # 判断主送机关（以冒号结尾且较短）
        if (line.endswith("：") or line.endswith(":")) and len(line) < 30 and not line.startswith("一、"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)  # 顶格
            r = p.add_run(line)
            r.font.name = tpl.body_font or "仿宋_GB2312"
            r.font.size = Pt(tpl.body_size_pt or 16.0)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), tpl.body_font or "仿宋_GB2312")
            continue

        # 判断成文日期 / 发文机关署名（位于末尾）
        if re.search(r"^\d{4}年\d{1,2}月\d{1,2}日$", line) or (len(line) < 18 and any(k in line for k in ("部", "局", "厅", "委员会", "办公室", "院", "处")) and not line.startswith("一、")):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.first_line_indent = Pt(0)
            r = p.add_run(line)
            r.font.name = tpl.body_font or "仿宋_GB2312"
            r.font.size = Pt(tpl.body_size_pt or 16.0)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), tpl.body_font or "仿宋_GB2312")
            continue

        # 首行缩进 2 字符 (32pt)
        p.paragraph_format.first_line_indent = Pt(32)

        # 一级标题：黑体
        if re.match(r"^[一二三四五六七八九十]+、", line):
            r = p.add_run(line)
            r.font.name = "黑体"
            r.font.size = Pt(tpl.body_size_pt or 16.0)
            r.font.bold = True
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        # 二级标题：楷体
        elif re.match(r"^（[一二三四五六七八九十]+）", line):
            r = p.add_run(line)
            r.font.name = "楷体_GB2312"
            r.font.size = Pt(tpl.body_size_pt or 16.0)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体_GB2312")
        # 正文普通段落：仿宋
        else:
            r = p.add_run(line)
            r.font.name = tpl.body_font or "仿宋_GB2312"
            r.font.size = Pt(tpl.body_size_pt or 16.0)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), tpl.body_font or "仿宋_GB2312")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
